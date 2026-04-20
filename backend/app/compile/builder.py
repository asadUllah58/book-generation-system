"""Compile approved chapters into .docx + .txt and upload to Supabase Storage.

Markdown support is intentionally minimal — we only recognise `#`/`##`/`###`
headings (emitted as Word headings) and render everything else as paragraphs.
If later we need richer markdown (bold, italic, lists), pipe through a proper
markdown parser; for a demo this is sufficient and has zero extra deps.
"""

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from app.db import repositories as repo
from app.db.supabase_client import get_supabase

BUCKET = "books"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
PDF_CONTENT_TYPE = "application/pdf"


# ---------- gather ----------

def approved_chapters_in_order(book_id: str) -> list[dict[str, Any]]:
    """Latest approved version per chapter index, sorted by index."""
    rows = repo.list_chapters(book_id)
    by_index: dict[int, dict[str, Any]] = {}
    for c in rows:
        if c["status"] != "approved":
            continue
        existing = by_index.get(c["index"])
        if not existing or c["version"] > existing["version"]:
            by_index[c["index"]] = c
    return [by_index[i] for i in sorted(by_index)]


# ---------- builders ----------

def build_docx(book_title: str, chapters: list[dict[str, Any]]) -> bytes:
    doc = Document()
    title_p = doc.add_heading(book_title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for ch in chapters:
        heading = ch.get("title") or f"Chapter {ch['index'] + 1}"
        doc.add_heading(heading, level=1)

        content_md = ch.get("content_md") or ""
        for raw in content_md.splitlines():
            line = raw.rstrip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                # The chapter heading was already added above.
                continue
            else:
                doc.add_paragraph(line)
        doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(book_title: str, chapters: list[dict[str, Any]]) -> bytes:
    """Render a PDF with a centred title page, per-chapter headings, and prose.

    Markdown handling matches `build_docx` — only `#`/`##`/`###` headings
    become real PDF headings; everything else renders as body paragraphs.
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=book_title,
    )
    base = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BookTitle", parent=base["Title"], fontSize=30, leading=36, spaceAfter=20, alignment=1
    )
    chapter_style = ParagraphStyle(
        "ChapterTitle", parent=base["Heading1"], fontSize=20, leading=24, spaceAfter=14
    )
    h2_style = ParagraphStyle(
        "H2", parent=base["Heading2"], fontSize=15, leading=18, spaceAfter=10
    )
    h3_style = ParagraphStyle(
        "H3", parent=base["Heading3"], fontSize=13, leading=16, spaceAfter=8
    )
    body_style = ParagraphStyle(
        "Body", parent=base["BodyText"], fontSize=11, leading=16, spaceAfter=6
    )

    def _escape(text: str) -> str:
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story: list = [Paragraph(_escape(book_title), title_style), PageBreak()]

    for ch in chapters:
        heading = ch.get("title") or f"Chapter {ch['index'] + 1}"
        story.append(Paragraph(_escape(heading), chapter_style))

        for raw in (ch.get("content_md") or "").splitlines():
            line = raw.rstrip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            if line.startswith("### "):
                story.append(Paragraph(_escape(line[4:].strip()), h3_style))
            elif line.startswith("## "):
                story.append(Paragraph(_escape(line[3:].strip()), h2_style))
            elif line.startswith("# "):
                # Chapter heading already added.
                continue
            else:
                story.append(Paragraph(_escape(line), body_style))
        story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


def build_txt(book_title: str, chapters: list[dict[str, Any]]) -> str:
    parts: list[str] = [book_title, "=" * len(book_title), ""]
    for ch in chapters:
        heading = ch.get("title") or f"Chapter {ch['index'] + 1}"
        parts.append(heading)
        parts.append("-" * len(heading))
        parts.append("")
        for line in (ch.get("content_md") or "").splitlines():
            if line.lstrip().startswith("#"):
                continue
            parts.append(line)
        parts.append("")
        parts.append("")
    return "\n".join(parts).rstrip() + "\n"


# ---------- storage ----------

def _upload(path: str, content: bytes, content_type: str) -> None:
    bucket = get_supabase().storage.from_(BUCKET)
    # supabase-py raises if the object already exists — remove first, then upload.
    try:
        bucket.remove([path])
    except Exception:
        pass
    bucket.upload(
        path=path,
        file=content,
        file_options={"content-type": content_type},
    )


def create_signed_url(path: str, expires_in_seconds: int = 3600) -> str:
    res = get_supabase().storage.from_(BUCKET).create_signed_url(
        path, expires_in_seconds
    )
    # supabase-py has shifted between `signedURL` and `signedUrl` across versions.
    return (
        res.get("signedURL")
        or res.get("signedUrl")
        or res.get("signed_url")
        or ""
    )


# ---------- entry point ----------

def compile_book(book_id: str, book_title: str) -> dict[str, str]:
    """Build and upload all three formats. Used by the graph's compile_draft
    node at end-of-pipeline."""
    chapters = approved_chapters_in_order(book_id)
    if not chapters:
        raise RuntimeError(f"no approved chapters to compile for book {book_id}")

    docx_bytes = build_docx(book_title, chapters)
    pdf_bytes = build_pdf(book_title, chapters)
    txt_bytes = build_txt(book_title, chapters).encode("utf-8")

    docx_path = f"{book_id}/book.docx"
    pdf_path = f"{book_id}/book.pdf"
    txt_path = f"{book_id}/book.txt"

    _upload(docx_path, docx_bytes, DOCX_CONTENT_TYPE)
    _upload(pdf_path, pdf_bytes, PDF_CONTENT_TYPE)
    _upload(txt_path, txt_bytes, "text/plain; charset=utf-8")

    return {"docx_path": docx_path, "pdf_path": pdf_path, "txt_path": txt_path}


def compile_book_format(book_id: str, book_title: str, fmt: str) -> str:
    """Build and upload a single format on-demand.

    Used by the download endpoint so reviewers can snapshot a book at any
    stage — even with only one approved chapter. Reuses the same Storage
    paths as `compile_book` so a later full compile just overwrites.
    """
    chapters = approved_chapters_in_order(book_id)
    if not chapters:
        raise RuntimeError(f"no approved chapters to compile for book {book_id}")

    if fmt == "docx":
        content = build_docx(book_title, chapters)
        path = f"{book_id}/book.docx"
        _upload(path, content, DOCX_CONTENT_TYPE)
    elif fmt == "pdf":
        content = build_pdf(book_title, chapters)
        path = f"{book_id}/book.pdf"
        _upload(path, content, PDF_CONTENT_TYPE)
    elif fmt == "txt":
        content = build_txt(book_title, chapters).encode("utf-8")
        path = f"{book_id}/book.txt"
        _upload(path, content, "text/plain; charset=utf-8")
    else:
        raise ValueError(f"unknown format: {fmt}")

    return path
