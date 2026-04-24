"""Compile outputs (combined book or single chapter) into docx / pdf / txt / md.

Orchestrates: gather chapters/outline → build bytes via format builders →
upload to Supabase Storage → return storage path for signed-URL issuance.

Markdown handling in format builders is intentionally minimal — `#`/`##`/`###`
heading lines become real headings, everything else is body prose.
"""

from io import BytesIO
from typing import Any, Literal, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from app.db import repositories as repo
from app.db.supabase_client import get_supabase


BUCKET = "books"
Format = Literal["docx", "pdf", "txt", "md"]

CONTENT_TYPES: dict[str, str] = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "txt": "text/plain; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
}


# ---------- gather ----------

def approved_chapters_in_order(book_id: str) -> list[dict[str, Any]]:
    rows = repo.list_chapters(book_id)
    by_index: dict[int, dict[str, Any]] = {}
    for c in rows:
        if c["status"] != "approved":
            continue
        existing = by_index.get(c["index"])
        if not existing or c["version"] > existing["version"]:
            by_index[c["index"]] = c
    return [by_index[i] for i in sorted(by_index)]


def latest_chapter_at(book_id: str, index: int) -> Optional[dict[str, Any]]:
    rows = [c for c in repo.list_chapters(book_id) if c["index"] == index]
    if not rows:
        return None
    return max(rows, key=lambda c: c["version"])


# ---------- docx ----------

def _docx_write_markdown(doc: Document, content_md: str) -> None:
    for raw in content_md.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            # Chapter heading is added separately by the caller.
            continue
        else:
            doc.add_paragraph(line)


def build_docx(
    book_title: str,
    chapters: list[dict[str, Any]],
    outline: Optional[dict[str, Any]] = None,
) -> bytes:
    doc = Document()
    title_p = doc.add_heading(book_title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    if outline:
        doc.add_heading("Outline", level=1)
        if outline.get("summary"):
            doc.add_paragraph(outline["summary"])
        for ch in outline.get("chapters", []):
            doc.add_heading(
                f"Chapter {ch['index'] + 1} — {ch['title']}", level=2
            )
            doc.add_paragraph(ch.get("summary", ""))
        doc.add_page_break()

    for ch in chapters:
        heading = ch.get("title") or f"Chapter {ch['index'] + 1}"
        doc.add_heading(heading, level=1)
        _docx_write_markdown(doc, ch.get("content_md") or "")
        doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- pdf ----------

def _pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "BookTitle", parent=base["Title"], fontSize=30, leading=36,
            spaceAfter=20, alignment=1,
        ),
        "h1": ParagraphStyle(
            "H1", parent=base["Heading1"], fontSize=20, leading=24, spaceAfter=14
        ),
        "h2": ParagraphStyle(
            "H2", parent=base["Heading2"], fontSize=15, leading=18, spaceAfter=10
        ),
        "h3": ParagraphStyle(
            "H3", parent=base["Heading3"], fontSize=13, leading=16, spaceAfter=8
        ),
        "body": ParagraphStyle(
            "Body", parent=base["BodyText"], fontSize=11, leading=16, spaceAfter=6
        ),
    }


def _pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_markdown_flow(content_md: str, styles: dict[str, ParagraphStyle]) -> list:
    story: list = []
    for raw in content_md.splitlines():
        line = raw.rstrip()
        if not line:
            story.append(Spacer(1, 6))
        elif line.startswith("### "):
            story.append(Paragraph(_pdf_escape(line[4:].strip()), styles["h3"]))
        elif line.startswith("## "):
            story.append(Paragraph(_pdf_escape(line[3:].strip()), styles["h2"]))
        elif line.startswith("# "):
            continue
        else:
            story.append(Paragraph(_pdf_escape(line), styles["body"]))
    return story


def build_pdf(
    book_title: str,
    chapters: list[dict[str, Any]],
    outline: Optional[dict[str, Any]] = None,
) -> bytes:
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
        title=book_title,
    )
    styles = _pdf_styles()
    story: list = [Paragraph(_pdf_escape(book_title), styles["title"]), PageBreak()]

    if outline:
        story.append(Paragraph("Outline", styles["h1"]))
        if outline.get("summary"):
            story.append(Paragraph(_pdf_escape(outline["summary"]), styles["body"]))
            story.append(Spacer(1, 10))
        for ch in outline.get("chapters", []):
            story.append(Paragraph(
                _pdf_escape(f"Chapter {ch['index'] + 1} — {ch['title']}"),
                styles["h2"],
            ))
            if ch.get("summary"):
                story.append(Paragraph(_pdf_escape(ch["summary"]), styles["body"]))
        story.append(PageBreak())

    for ch in chapters:
        heading = ch.get("title") or f"Chapter {ch['index'] + 1}"
        story.append(Paragraph(_pdf_escape(heading), styles["h1"]))
        story.extend(_pdf_markdown_flow(ch.get("content_md") or "", styles))
        story.append(PageBreak())

    doc.build(story)
    return buf.getvalue()


# ---------- txt ----------

def build_txt(
    book_title: str,
    chapters: list[dict[str, Any]],
    outline: Optional[dict[str, Any]] = None,
) -> str:
    parts: list[str] = [book_title, "=" * len(book_title), ""]

    if outline:
        parts.extend(["Outline", "-------", ""])
        if outline.get("summary"):
            parts.extend([outline["summary"], ""])
        for ch in outline.get("chapters", []):
            parts.append(f"Chapter {ch['index'] + 1} — {ch['title']}")
            if ch.get("summary"):
                parts.append(f"  {ch['summary']}")
            parts.append("")
        parts.append("")

    for ch in chapters:
        heading = ch.get("title") or f"Chapter {ch['index'] + 1}"
        parts.append(heading)
        parts.append("-" * len(heading))
        parts.append("")
        for line in (ch.get("content_md") or "").splitlines():
            if line.lstrip().startswith("#"):
                continue
            parts.append(line)
        parts.append("")
        parts.append("")
    return "\n".join(parts).rstrip() + "\n"


# ---------- md ----------

def build_md(
    book_title: str,
    chapters: list[dict[str, Any]],
    outline: Optional[dict[str, Any]] = None,
) -> str:
    parts: list[str] = [f"# {book_title}", ""]

    if outline:
        parts.extend(["## Outline", ""])
        if outline.get("summary"):
            parts.extend([outline["summary"], ""])
        for ch in outline.get("chapters", []):
            parts.append(f"### Chapter {ch['index'] + 1} — {ch['title']}")
            if ch.get("summary"):
                parts.append("")
                parts.append(ch["summary"])
            parts.append("")

    for ch in chapters:
        heading = ch.get("title") or f"Chapter {ch['index'] + 1}"
        parts.append(f"## {heading}")
        parts.append("")
        parts.append(ch.get("content_md") or "")
        parts.append("")
    return "\n".join(parts).rstrip() + "\n"


# ---------- storage ----------

def _upload(path: str, content: bytes, content_type: str) -> None:
    bucket = get_supabase().storage.from_(BUCKET)
    try:
        bucket.remove([path])
    except Exception:
        pass
    bucket.upload(
        path=path,
        file=content,
        file_options={"content-type": content_type},
    )


def create_signed_url(path: str, expires_in_seconds: int = 3600) -> str:
    res = get_supabase().storage.from_(BUCKET).create_signed_url(
        path, expires_in_seconds
    )
    return (
        res.get("signedURL")
        or res.get("signedUrl")
        or res.get("signed_url")
        or ""
    )


# ---------- entry points ----------

def _encode(text: str) -> bytes:
    return text.encode("utf-8")


def _build(
    fmt: Format,
    book_title: str,
    chapters: list[dict[str, Any]],
    outline: Optional[dict[str, Any]],
) -> bytes:
    if fmt == "docx":
        return build_docx(book_title, chapters, outline)
    if fmt == "pdf":
        return build_pdf(book_title, chapters, outline)
    if fmt == "txt":
        return _encode(build_txt(book_title, chapters, outline))
    if fmt == "md":
        return _encode(build_md(book_title, chapters, outline))
    raise ValueError(f"unknown format: {fmt}")


def compile_combined(book_id: str, fmt: Format) -> str:
    """Build the combined book (outline + approved chapters) and upload it.

    Returns the Storage path. Same path is reused across invocations so a
    later call simply overwrites the previous snapshot.
    """
    book = repo.get_book(book_id)
    if not book:
        raise LookupError(f"book {book_id} not found")

    chapters = approved_chapters_in_order(book_id)
    if not chapters:
        raise RuntimeError("no approved chapters to compile")

    outline_row = repo.get_approved_outline(book_id)
    outline_content = outline_row["content"] if outline_row else None

    content = _build(fmt, book["title"], chapters, outline_content)
    path = f"{book_id}/book.{fmt}"
    _upload(path, content, CONTENT_TYPES[fmt])
    return path


def compile_chapter(book_id: str, index: int, fmt: Format) -> str:
    """Build a single-chapter file and upload it."""
    book = repo.get_book(book_id)
    if not book:
        raise LookupError(f"book {book_id} not found")

    chapter = latest_chapter_at(book_id, index)
    if not chapter or not chapter.get("content_md"):
        raise RuntimeError(f"chapter {index} has no content to download")

    # Outline is intentionally omitted for per-chapter downloads.
    content = _build(fmt, book["title"], [chapter], None)
    path = f"{book_id}/chapter-{index}.{fmt}"
    _upload(path, content, CONTENT_TYPES[fmt])
    return path
